# utils.py
import streamlit as st
from pydrive2.auth import GoogleAuth
from pydrive2.drive import GoogleDrive
from oauth2client.service_account import ServiceAccountCredentials
from validate_docbr import CPF, CNPJ
import json # json ainda é necessário para o arquivo local
import requests
from datetime import date, datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import io
import uuid

# --- FUNÇÃO DE LOGIN COM CONTA DE SERVIÇO ---
def login_gdrive():
    gauth = GoogleAuth()
    scope = ["https://www.googleapis.com/auth/drive"]
    
    if hasattr(st, 'secrets') and 'gdrive_service_account' in st.secrets:
        # AQUI st.secrets["gdrive_service_account"] DEVE SER UM DICIONÁRIO
        # Se este erro persiste, o problema está na forma como o secret foi COLADO no Streamlit Cloud.
        creds_dict = st.secrets["gdrive_service_account"] 
        gauth.credentials = ServiceAccountCredentials.from_json_keyfile_dict(creds_dict, scope)
    else:
        # Fallback para arquivo local 'service_account.json' para desenvolvimento local
        try:
            with open('service_account.json', 'r') as f:
                creds_dict = json.load(f) # Aqui sim, precisa carregar o JSON do arquivo
            gauth.credentials = ServiceAccountCredentials.from_json_keyfile_dict(creds_dict, scope)
        except FileNotFoundError:
            st.error("Arquivo de autenticação local 'service_account.json' não encontrado. Para deploy, configure st.secrets.")
            st.stop()
        except Exception as e:
            st.error(f"Erro ao carregar credenciais: {e}. Para deploy, configure st.secrets.")
            st.stop()

    return GoogleDrive(gauth)

# --- FUNÇÃO PARA GERAR O CONTRATO EM WORD (COM ALTERAÇÕES DE FORMATAÇÃO) ---
def gerar_contrato_docx(dados):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    p_format = style.paragraph_format
    p_format.line_spacing = 1.5
    p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    section = doc.sections[0]
    header = section.header
    p_header = header.paragraphs[0]
    run_header = p_header.add_run()
    try:
        run_header.add_picture('assets/logo.png', width=Inches(2.0))
    except FileNotFoundError:
        p_header.text = "Rocker Equipamentos"
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = titulo.add_run(f"CONTRATO DE {dados['tipo_contrato'].upper()} Nº {dados['numero_contrato']}\n")
    run_titulo.bold = True
    run_titulo.font.size = Pt(14)

    def add_justified_paragraph(text=''):
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p

    def add_clausula_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        return p
    
    add_clausula_heading("DAS PARTES")
    doc.add_paragraph("")
    
    p_locadora = add_justified_paragraph()
    p_locadora.add_run("LOCADORA: ").bold = True
    p_locadora.add_run("ROCKER LOCAÇÃO DE EQUIPAMENTOS PARA CONSTRUÇÃO LTDA, pessoa jurídica de direito privado, inscrita no CNPJ sob o nº 15.413.157/0001-16, com sede na Rua Carlos Adriano Rodrigues da Silva, Q40 L01, Bairro Potecas, São José/SC, CEP 88.107-493, neste ato representada na forma de seu contrato social.")
    
    p_locataria = add_justified_paragraph()
    p_locataria.add_run("LOCATÁRIA: ").bold = True
    if dados['cliente']['tipo_pessoa'] == "Pessoa Jurídica":
        p_locataria.add_run(f"{dados['cliente']['nome_razao_social']}, pessoa jurídica de direito privado, inscrita no CNPJ sob o nº {dados['cliente']['cpf_cnpj']}, com sede na {dados['cliente']['endereco']}, {dados['cliente']['cidade']} - {dados['cliente']['estado']}, CEP: {dados['cliente']['cep']}, neste ato representada por seu representante legal, {dados['cliente']['representante_legal']['nome']}, portador(a) do CPF sob o nº {dados['cliente']['representante_legal']['cpf']}.")
    else:
        p_locataria.add_run(f"{dados['cliente']['nome_razao_social']}, inscrito(a) no CPF sob o nº {dados['cliente']['cpf_cnpj']}, residente e domiciliado(a) na {dados['cliente']['endereco']}, {dados['cliente']['cidade']} - {dados['cliente']['estado']}, CEP: {dados['cliente']['cep']}.")
    
    add_justified_paragraph("\nAs partes acima qualificadas celebram o presente contrato, que se regerá pelas cláusulas e condições a seguir.")
    
    add_clausula_heading("\nCLÁUSULA PRIMEIRA – DO OBJETO")
    add_justified_paragraph("1.1. O objeto deste contrato é a locação do(s) equipamento(s) descrito(s) na Cláusula Segunda, para ser(em) utilizado(s) exclusivamente no endereço da obra informado abaixo.")

    add_clausula_heading("\nCLÁUSULA SEGUNDA – DOS EQUIPAMENTOS, VALORES E CONDIÇÕES")
    add_justified_paragraph("2.1. Equipamentos e Valores da Locação:")

    tabela = doc.add_table(rows=1, cols=5)
    tabela.style = 'Table Grid'
    tabela.autofit = False
    tabela.allow_autofit = False
    
    tabela.columns[0].width = Cm(1.3)
    tabela.columns[1].width = Cm(1.3)
    tabela.columns[2].width = Cm(7.0)
    tabela.columns[3].width = Cm(3.0)
    tabela.columns[4].width = Cm(3.0)

    hdr_cells = tabela.rows[0].cells
    hdr_cells[0].text = 'Item'
    hdr_cells[1].text = 'Qtde'
    hdr_cells[2].text = 'Equipamento'
    hdr_cells[3].text = 'Vlr. Unit. Mensal (R$)'
    hdr_cells[4].text = 'Vlr. Total Mensal (R$)'

    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
    total_locacao_mensal = 0
    for i, item in enumerate(dados['itens_contrato']):
        row_cells = tabela.add_row().cells
        row_cells[0].text = f"2.1.{i+1}"
        row_cells[1].text = str(item['quantidade'])
        row_cells[2].text = f"{item['produto']} COM {item['plataforma']}"
        row_cells[3].text = f"{item['valor_unitario']:.2f}"
        valor_total_item = item['quantidade'] * item['valor_unitario']
        row_cells[4].text = f"{valor_total_item:.2f}"
        total_locacao_mensal += valor_total_item
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
    add_justified_paragraph("\n2.2. Resumo Financeiro:")
    add_justified_paragraph(f"Valor Total da Locação Mensal: R$ {total_locacao_mensal:.2f}")
    add_justified_paragraph(f"Custo de Entrega (Frete): R$ {dados['valor_entrega']:.2f}")
    add_justified_paragraph(f"Custo de Recolha (Frete): R$ {dados['valor_recolha']:.2f}")

    add_justified_paragraph("\n2.3. Contato e Endereço da Obra:")
    add_justified_paragraph(f"Contato Responsável na Obra: {dados['contato_nome']}")
    add_justified_paragraph(f"Telefone: {dados['contato_telefone']}")
    add_justified_paragraph(f"Endereço da Obra: {dados['endereco_obra']}")

    add_clausula_heading("\nCLÁUSULA TERCEIRA – DO PRAZO")
    add_justified_paragraph(f"3.1. A locação terá início em {dados['data_inicio']} e se encerrará com a devolução integral dos equipamentos à LOCADORA em perfeitas condições de uso. Para equipamentos não devolvidos ou danificados, a locação permanecerá vigente até a quitação da indenização correspondente.")
    add_justified_paragraph("3.2. A vigência mínima deste contrato é de 30 (trinta) dias. Após este período, caso não haja manifestação de rescisão por escrito, o contrato será renovado automaticamente por períodos iguais e sucessivos.")
    add_justified_paragraph("3.3. Caso a LOCATÁRIA não comunique por escrito a intenção de devolver os equipamentos com antecedência mínima de 05 (cinco) dias do vencimento do período vigente, a LOCADORA fica autorizada a faturar um novo período de 30 (trinta) dias.")

    add_clausula_heading("\nCLÁUSULA QUARTA – DO FATURAMENTO")
    add_justified_paragraph("4.1. O faturamento do primeiro período de 30 (trinta) dias será realizado no ato da entrega dos equipamentos. Os períodos subsequentes serão faturados a cada 30 (trinta) dias.")
    
    add_clausula_heading("\nCLÁUSULA QUINTA – DO PAGAMENTO")
    add_justified_paragraph("5.1. O pagamento do primeiro período será efetuado na data da entrega dos equipamentos. Os pagamentos subsequentes deverão ser realizados a cada 30 (trinta) dias.")
    add_justified_paragraph("5.2. Os pagamentos serão realizados em moeda corrente nacional, por meio de boleto bancário, que a LOCATÁRIA desde já reconhece como dívida líquida, certa e exigível em seu vencimento.")
    add_justified_paragraph("5.3. O atraso no pagamento sujeitará a LOCATÁRIA à multa de 3% (três por cento) sobre o valor do débito, acrescido de juros e correção monetária. Após 5 (cinco) dias de atraso, o título poderá ser enviado a protesto.")

    add_clausula_heading("\nCLÁUSULA SEXTA – DO LOCAL E USO DOS EQUIPAMENTOS")
    add_justified_paragraph("6.1. Os equipamentos deverão ser utilizados exclusivamente no endereço da obra informado na Cláusula Segunda.")
    add_justified_paragraph("6.2. É expressamente vedado à LOCATÁRIA ceder, emprestar, sublocar ou transferir os equipamentos a terceiros ou para outro local, sob qualquer pretexto.")
    
    add_clausula_heading("\nCLÁUSULA SÉTIMA – DA GUARDA E CONSERVAÇÃO")
    add_justified_paragraph("7.1. Ao assinar o check list de entrega, a LOCATÁRIA declara ter recebido todos os componentes do equipamento objeto deste contrato em perfeito estado de conservação e uso.")
    add_justified_paragraph("7.2. A LOCATÁRIA responsabiliza-se inteiramente pela guarda e conservação dos equipamentos de propriedade da LOCADORA. Eventuais danos por mau uso ou imperícia ou falta de todo ou partes dos equipamentos serão indenizados pela LOCATÁRIA.")
    add_justified_paragraph("7.3. Será considerado o valor de indenização para cada componente, o mesmo constante no check list de entrega.")

    add_clausula_heading("\nCLÁUSULA OITAVA – DA UTILIZAÇÃO")
    add_justified_paragraph("8.1. A montagem, desmontagem e correta utilização dos equipamentos ficarão exclusivamente, a cargo da LOCATÁRIA. Caso a LOCATÁRIA queira assistência da LOCADORA, devera solicitar isso por escrito.")
    add_justified_paragraph("8.2. A LOCADORA não se responsabiliza por eventuais acidentes decorrentes do mau uso, erro ou falha na operação e negligência em relação às normas de segurança do trabalho durante a vigência deste contrato, o mesmo será de responsabilidade da LOCATÁRIA, sejam civis, criminais ou trabalhistas.")
    add_justified_paragraph("8.3. A LOCADORA envia junto do equipamento placas de identificação e capacidade de carga de acordo com o tamanho de cada balancim suspenso locado, sendo responsabilidade da LOCATÁRIA a fixação das mesmas antes da utilização do equipamento.")

    add_clausula_heading("\nCLÁUSULA NONA – DA MANUTENÇÃO")
    add_justified_paragraph("9.1. A LOCATÁRIA fica obrigada a zelar pela boa manutenção do equipamento, mantendo sempre todas as peças e partes em funcionamento. Fica ainda sob responsabilidade da LOCATÁRIA as manutenções preventivas que se fizerem necessárias durante a vigência deste contrato.")
    add_justified_paragraph("9.2. Para reparos que necessitem mão de obra especializada, a LOCATÁRIA deverá solicitar por escrito a assistência técnica da LOCADORA, que deverá atender a solicitação em um prazo de 48 horas após a comunicação.")
    add_justified_paragraph("9.3. Caso seja constatado que a danificação apresentada no equipamento seja por mau uso, erro ou falha na operação, vandalismo e/ou falta de energia elétrica, o custo de manutenção e reparos ocorrerá por conta da LOCATÁRIA.")

    add_clausula_heading("\nCLÁUSULA DÉCIMA – DO DIREITO DE PROPRIEDADE")
    add_justified_paragraph("10.1. Fica resguardado o direito de propriedade da LOCADORA sobre o equipamento contratado, acima de qualquer situação, condição ou pretexto alegados pela LOCATÁRIA ou por terceiros.")
    
    add_clausula_heading("\nCLÁUSULA DÉCIMA PRIMEIRA – DA RESCISÃO")
    p_rescisao_intro = doc.add_paragraph("11.1. O presente contrato fica rescindido de pleno direito, independente de qualquer aviso, sem prejuízo das penalidades, nas seguintes hipóteses:")
    p_rescisao_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph("")

    p_item_I = doc.add_paragraph("I - Atraso no pagamento de 2 (dois) aluguéis consecutivos;")
    p_item_I.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph("")

    p_item_II = doc.add_paragraph("II - Infração das demais cláusulas contratuais;")
    p_item_II.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph("")

    p_item_III = doc.add_paragraph("III - Falência, insolvência ou concordata da LOCATÁRIA;")
    p_item_III.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph("")

    p_item_IV = doc.add_paragraph("IV - Quando a LOCADORA achar e julgar necessária o encerramento do contrato.")
    p_item_IV.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph("")

    add_clausula_heading("CLÁUSULA DÉCIMA SEGUNDA – DO FORO")
    add_justified_paragraph("12.1. Fica eleito o foro central da comarca de São José para dirimir eventuais litígios oriundos deste contrato, se solução amigável não advir.")

    p_final_intro = doc.add_paragraph("E, por estarem justas e contratadas, as partes firmam o presente instrumento em 2 (duas) vias de igual teor e forma, na presença das duas testemunhas abaixo.")
    p_final_intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    meses_pt = {
        "january": "janeiro", "february": "fevereiro", "march": "março",
        "april": "abril", "may": "maio", "june": "junho", "july": "julho",
        "august": "agosto", "september": "setembro", "october": "outubro",
        "november": "novembro", "december": "dezembro"
    }
    
    data_assinatura_str = dados['data_assinatura']
    partes_data = data_assinatura_str.split(' de ')
    data_formatada_pt = data_assinatura_str
    if len(partes_data) == 3:
        dia = partes_data[0]
        mes_ingles = partes_data[1]
        ano = partes_data[2]
        mes_portugues = meses_pt.get(mes_ingles.lower(), mes_ingles)
        data_formatada_pt = f"{dia} de {mes_portugues} de {ano}"

    assinatura_data = doc.add_paragraph(f"\nSão José, {data_formatada_pt}.")
    assinatura_data.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p_assinatura_locadora_linha = doc.add_paragraph("_________________________________________")
    p_assinatura_locadora_linha.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_assinatura_locadora_nome = doc.add_paragraph("ROCKER LOCAÇÃO DE EQUIPAMENTOS LTDA")
    p_assinatura_locadora_nome.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_assinatura_locadora_qualif = doc.add_paragraph("(LOCADORA)")
    p_assinatura_locadora_qualif.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph("")
    
    p_assinatura_locataria_linha = doc.add_paragraph("_________________________________________")
    p_assinatura_locataria_linha.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_assinatura_locataria_nome = doc.add_paragraph(f"{dados['cliente']['nome_razao_social'].upper()}")
    p_assinatura_locataria_nome.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_assinatura_locataria_qualif = doc.add_paragraph("(LOCATÁRIA)")
    p_assinatura_locataria_qualif.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph("")
    
    p_testemunhas_titulo = doc.add_paragraph("Testemunhas:")
    p_testemunhas_titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("")

    p_testemunha1_linha = doc.add_paragraph("_________________________________________")
    p_testemunha1_linha.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_testemunha1_nome = doc.add_paragraph("Nome:")
    p_testemunha1_nome.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_testemunha1_cpf = doc.add_paragraph("CPF:")
    p_testemunha1_cpf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("")
    
    p_testemunha2_linha = doc.add_paragraph("_________________________________________")
    p_testemunha2_linha.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_testemunha2_nome = doc.add_paragraph("Nome:")
    p_testemunha2_nome.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_testemunha2_cpf = doc.add_paragraph("CPF:")
    p_testemunha2_cpf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- FUNÇÃO PARA GERAR A FATURA EM WORD (COM CORREÇÕES E MELHORIAS) ---
def gerar_fatura_docx(dados_fatura):
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_logo = p_logo.add_run()
    try:
        run_logo.add_picture('assets/logo.png', width=Cm(14.52), height=Cm(2.22)) 
    except FileNotFoundError:
        run_logo.text = "Rocker Equipamentos"
    
    tabela_principal = doc.add_table(rows=7, cols=2)
    tabela_principal.style = 'Table Grid'
    tabela_principal.autofit = False
    tabela_principal.allow_autofit = False
    
    tabela_principal.columns[0].width = Cm(10.0)
    tabela_principal.columns[1].width = Cm(7.0)

    celula_rocker = tabela_principal.cell(0, 0)
    celula_fatura = tabela_principal.cell(0, 1)
    
    celula_rocker.paragraphs[0].clear()
    p_rocker_info = celula_rocker.add_paragraph()
    run_rocker_name = p_rocker_info.add_run("ROCKER LOCAÇÃO DE EQUP. PARA CONST. LTDA EPP\n")
    run_rocker_name.bold = True
    p_rocker_info.add_run("Rua Carlos Adriano Rodrigues da Silva, S/N – Potecas – São José SC\n"
                            "CNPJ 15.413.157/0001-16\n"
                            "Telefones: 48 3246.9987 - 3372.6439")
    
    celula_fatura.paragraphs[0].clear()
    p_titulo = celula_fatura.add_paragraph("FATURA DE LOCAÇÃO")
    p_titulo.runs[0].bold = True
    celula_fatura.add_paragraph(f"N°{dados_fatura['NUMERO_FATURA']}")
    celula_fatura.add_paragraph(f"Emissão: {dados_fatura['DATA_EMISSAO']}")

    celula_dest_titulo = tabela_principal.cell(1, 0).merge(tabela_principal.cell(1, 1))
    celula_dest_titulo.text = "DESTINATÁRIO"
    celula_dest_titulo.paragraphs[0].runs[0].bold = True

    celula_cliente_info_end = tabela_principal.cell(2, 0).merge(tabela_principal.cell(2, 1))
    
    p_cliente = celula_cliente_info_end.paragraphs[0] if celula_cliente_info_end.paragraphs else celula_cliente_info_end.add_paragraph()
    p_cliente.clear()
    p_cliente.add_run(f"{dados_fatura['NOME_CLIENTE']}\nCNPJ/CPF: {dados_fatura['CNPJ_CLIENTE']}\n")
    
    run_endereco_label = p_cliente.add_run("Endereço: ")
    run_endereco_label.bold = True
    
    endereco_completo = f"{dados_fatura['ENDERECO_CLIENTE']}, {dados_fatura.get('BAIRRO_CLIENTE', '')} - {dados_fatura.get('CIDADE_CLIENTE', '')} - {dados_fatura.get('ESTADO_CLIENTE', 'SC')}, CEP: {dados_fatura.get('CEP_CLIENTE', '')}"
    p_cliente.add_run(endereco_completo)
    
    celula_desc = tabela_principal.cell(3, 0).merge(tabela_principal.cell(3, 1))
    p_desc = celula_desc.paragraphs[0] if celula_desc.paragraphs else celula_desc.add_paragraph()
    p_desc.clear()
    run_desc_label = p_desc.add_run("Descrição: ")
    run_desc_label.bold = True
    p_desc.add_run(f"{dados_fatura['DESCRICAO_SERVICO']}")

    celula_valor = tabela_principal.cell(4, 0)
    celula_vencimento = tabela_principal.cell(4, 1)

    celula_valor.paragraphs[0].clear()
    p_valor = celula_valor.add_paragraph()
    run_valor_label = p_valor.add_run("Valor Total: ")
    run_valor_label.bold = True
    p_valor.add_run(f"R$ {float(dados_fatura['VALOR_TOTAL'].replace(',', '.')):.2f}")

    celula_vencimento.paragraphs[0].clear()
    p_venc = celula_vencimento.add_paragraph()
    run_venc_label = p_venc.add_run("Vencimento: ")
    run_venc_label.bold = True
    p_venc.add_run(f"{dados_fatura['DATA_VENCIMENTO']}")
    
    celula_forma_pag = tabela_principal.cell(5, 0).merge(tabela_principal.cell(5, 1))
    p_forma_pag = celula_forma_pag.paragraphs[0] if celula_forma_pag.paragraphs else celula_forma_pag.add_paragraph()
    p_forma_pag.clear()
    run_forma_pag_label = p_forma_pag.add_run("Forma de Pagamento: ")
    run_forma_pag_label.bold = True
    p_forma_pag.add_run(f"{dados_fatura['FORMA_PAGAMENTO']}")

    celula_obs = tabela_principal.cell(6, 0).merge(tabela_principal.cell(6, 1))
    p_obs = celula_obs.paragraphs[0] if celula_obs.paragraphs else celula_obs.add_paragraph()
    p_obs.clear()
    run_obs_label = p_obs.add_run("Observações: ")
    run_obs_label.bold = True
    
    observacao = dados_fatura.get('OBSERVACAO', '') 
    if observacao:
        run_obs_text = p_obs.add_run(observacao)
        run_obs_text.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    section = doc.sections[0]
    footer = section.footer
    p_footer = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p_footer.text = (
        "OPERAÇÃO NÃO TRIBUTADA PELO ISS, CONFORME LEI COMPL. Nº 116 DE 31/07/03. "
        "O serviço de locação de bens móveis não se enquadra no rol de serviço que trata do artigo 30 da lei 10833/03 não estando sujeito a retenção. "
        "Dispensa de emissão de Nota Fiscal de Serviços conforme Lei Complementar nº 116 de 31 de Julho de 2010."
    )
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_footer.runs[0].font.size = Pt(8)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- FUNÇÕES DE GERENCIAMENTO DE NÚMEROS (SEM ALTERAÇÕES) ---
def get_next_contract_number(drive):
    try:
        config_file = get_database_file(drive, "config.json")
        config_data = read_data(config_file)
        ultimo_numero = config_data.get("ultimo_numero_contrato", 0)
        novo_numero = ultimo_numero + 1
        ano_atual = date.today().year
        config_data["ultimo_numero_contrato"] = novo_numero
        write_data(config_file, config_data)
        return f"{str(novo_numero).zfill(5)}-{ano_atual}"
    except Exception as e:
        st.error(f"Erro ao obter número do contrato: {e}")
        return None

def get_next_fatura_number(drive):
    try:
        config_file = get_database_file(drive, "config.json")
        config_data = read_data(config_file)
        ultimo_numero = config_data.get("ultimo_numero_fatura", 0)
        novo_numero = ultimo_numero + 1
        config_data["ultimo_numero_fatura"] = novo_numero
        write_data(config_file, config_data)
        return f"{str(novo_numero).zfill(7)}"
    except Exception as e:
        st.error(f"Erro ao obter número da fatura: {e}")
        return None

# --- FUNÇÕES DE CONEXÃO COM GOOGLE DRIVE (SEM ALTERAÇÕES) ---
def get_database_file(drive, filename):
    file_list = drive.ListFile({'q': f"title='{filename}' and trashed=false"}).GetList()
    if file_list:
        return file_list[0]
    else:
        file = drive.CreateFile({'title': filename, 'mimeType': 'application/json'})
        file.SetContentString('[]')
        file.Upload()
        return file

def read_data(drive_file):
    content = drive_file.GetContentString()
    if not content:
        return []
    return json.loads(content)

def write_data(drive_file, data):
    drive_file.SetContentString(json.dumps(data, indent=4, ensure_ascii=False))
    drive_file.Upload()

# --- FUNÇÕES DE VALIDAÇÃO (SEM ALTERAÇÕES) ---
def validar_e_formatar_cpf(cpf_str):
    cpf = CPF()
    if cpf.validate(cpf_str):
        return cpf.mask(cpf_str)
    return None

def validar_e_formatar_cnpj(cnpj_str):
    cnpj = CNPJ()
    if cnpj.validate(cnpj_str):
        return cnpj.mask(cnpj_str)
    return None

# --- FUNÇÃO DE CONSULTA DE CEP (SEM ALTERAÇÕES) ---
def consultar_cep(cep):
    cep_limpo = "".join(filter(str.isdigit, cep))
    if len(cep_limpo) != 8:
        return None
    try:
        url = f"https://viacep.com.br/ws/{cep_limpo}/json/"
        response = requests.get(url)
        if response.status_code == 200:
            dados = response.json()
            if dados.get("erro"):
                return None
            return dados
        else:
            return None
    except requests.RequestException:
        return None

# --- COMPONENTE DE RODAPÉ (SEM ALTERAÇÕES) ---
def exibir_rodape():
    st.markdown("---")
    st.markdown(
        """
        <div style="text-align: center; font-size: 14px;">
            <p>© 2025 Gerenciamento de Clientes. Todos os direitos reservados.</p>
            <p>Desenvolvido com a expertise da 
                <a href="https://ascendtechdigital.com.br/" target="_blank">AscendTech</a>.
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
